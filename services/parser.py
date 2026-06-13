import os
import re
import pdfplumber
from pypdf import PdfReader
import docx

def clean_text(text):
    """
    Cleans extracted text by normalizing whitespaces, correcting common encoding issues,
    and stripping leading/trailing whitespace.
    """
    if not text:
        return ""
    # Replace non-breaking spaces with standard space
    text = text.replace('\xa0', ' ')
    # Normalize multiple consecutive spaces, tabs, and newlines to single characters
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n+', '\n', text)
    # Strip carriage returns
    text = text.replace('\r', '')
    return text.strip()

def parse_pdf(file_path):
    """
    Extracts text from a PDF file. Uses pdfplumber first for high quality tabular
    text reconstruction, and falls back to pypdf in case of failure.
    """
    text = ""
    try:
        # Attempt extraction with pdfplumber
        with pdfplumber.open(file_path) as pdf:
            pages = []
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
            text = "\n".join(pages)
    except Exception as e:
        print(f"[Parser] pdfplumber failed on {file_path}: {str(e)}. Falling back to pypdf.")
        # Fallback to pypdf
        try:
            reader = PdfReader(file_path)
            pages = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
            text = "\n".join(pages)
        except Exception as pypdf_err:
            print(f"[Parser] pypdf failed on {file_path}: {str(pypdf_err)}")
            raise Exception(f"Failed to parse PDF document: {str(pypdf_err)}")
            
    return clean_text(text)

def parse_docx(file_path):
    """
    Extracts text from a Word document (.docx). Iterates through paragraphs
    and table structures to assemble a coherent representation.
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Read normal paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # Read text from tables to avoid losing details inside tabular CV layouts
        for table in doc.tables:
            for row in table.rows:
                # Deduplicate cell values if cells are merged
                seen_cells = set()
                row_cells_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in seen_cells:
                        seen_cells.add(cell_text)
                        row_cells_text.append(cell_text)
                if row_cells_text:
                    full_text.append(" | ".join(row_cells_text))
                    
        return clean_text("\n".join(full_text))
    except Exception as e:
        print(f"[Parser] docx parsing failed: {str(e)}")
        raise Exception(f"Failed to parse DOCX document: {str(e)}")

def extract_text_from_file(file_path):
    """
    Master function to route file parsing based on extension.
    Supports PDF and DOCX/DOC.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    _, ext = os.path.splitext(file_path.lower())
    if ext == '.pdf':
        return parse_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
